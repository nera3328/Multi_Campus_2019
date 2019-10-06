import docx
import word
docName='sample-one-line.docx'
print('Document in full:\n', word.getTextWord(docName))

doc=docx.Document(docName)
print('단락 수 : ' , len(doc.paragraphs))
print('2번 단락 : ', doc.paragraphs[1].text)
print('2번 단락 스타일 : ', doc.paragraphs[1].style)

print('paragraph1:',doc.paragraphs[0].text)
print('number 1:',len(doc.paragraphs[0].runs))
for idx, run in enumerate(doc.paragraphs[0].runs):
   print('Run %s : %s' %(idx, run.text)) 


import os
import word,pdf
from nltk.corpus.reader.plaintext import PlaintextCorpusReader

def getText(txtFileName):
    file=open(txtFileName, 'r')
    return file.read()

newCorpusDir='mycorpus/'
if not os.path.isdir(newCorpusDir):
    os.mkdir(newCorpusDir)
    
txt1=getText('sample_feed.txt')
txt2=pdf.getTextPDF('sample-pdf.pdf')
txt3=word.getTextWord('sample-one-line.docx')

files=[txt1,txt2,txt3]
for idx, f in enumerate(files):
    with open(newCorpusDir+str(idx)+'.txt','w')as fout:
        fout.write(f)
newCorpus=PlaintextCorpusReader(newCorpusDir,'.*')
print(newCorpus.words())
print(newCorpus.sents(newCorpus.fileids()[1]))
print(newCorpus.paras(newCorpus.fileids()[0]))
